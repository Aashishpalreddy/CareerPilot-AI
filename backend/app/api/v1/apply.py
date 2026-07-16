import io
import logging
from pathlib import Path

from docx import Document
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from backend.app.api.dependencies import (
    get_daily_pipeline_service,
    get_saved_job_repository,
)
from backend.app.core.security import get_current_user
from backend.app.database.session import get_db
from backend.app.models.user import User
from backend.app.repositories.job_repository import JobRepository
from backend.app.repositories.saved_job_repository import SavedJobRepository
from backend.app.schemas.saved_job import SavedJobResponse
from backend.app.services.ai.daily_pipeline_service import (
    DailyPipelineService,
    TailoringNotPossible,
)

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/apply", tags=["Apply"])


class RunPipelineRequest(BaseModel):
    keywords: list[str]
    location: str | None = None
    remote_only: bool = False
    job_type: str | None = None          # full-time, contract, part-time
    experience_level: str | None = None  # entry, mid, senior
    work_arrangement: str | None = None  # onsite, hybrid, remote


@router.post("/run", response_model=list[SavedJobResponse])
async def run_daily_pipeline(
    request: RunPipelineRequest,
    current_user: User = Depends(get_current_user),
    pipeline: DailyPipelineService = Depends(get_daily_pipeline_service),
):
    """
    Runs discovery + matching + tailoring for the current user right now.
    In production this same call is what the daily scheduler triggers.
    """

    # Map work_arrangement to remote_only for backward compatibility
    remote_only = request.remote_only
    if request.work_arrangement == "remote":
        remote_only = True

    return await pipeline.run_for_user(
        user_id=current_user.id,
        keywords=request.keywords,
        location=request.location,
        remote_only=remote_only,
        job_type=request.job_type,
        experience_level=request.experience_level,
    )


@router.post("/process-job/{job_id}", response_model=SavedJobResponse)
async def process_job(
    job_id: int,
    current_user: User = Depends(get_current_user),
    pipeline: DailyPipelineService = Depends(get_daily_pipeline_service),
    db: Session = Depends(get_db),
):
    """
    On-demand tailoring for a single tracked job: matches it against the
    user's default resume, generates a tailored resume + cover letter, and
    auto-applies if the job is on a supported ATS. This is the per-job
    counterpart to /apply/run — search stays fast because this heavier
    work only runs for a job the user actually picks.
    """

    job = JobRepository(db).get_by_id(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")

    try:
        saved_job = await pipeline.process_job_for_user(
            user_id=current_user.id,
            job_id=job_id,
        )
    except TailoringNotPossible as e:
        raise HTTPException(status_code=422, detail=e.reason)

    return saved_job


@router.get("/saved-jobs", response_model=list[SavedJobResponse])
def list_saved_jobs(
    status_filter: str | None = None,
    current_user: User = Depends(get_current_user),
    repository: SavedJobRepository = Depends(get_saved_job_repository),
):
    return repository.get_by_user(
        user_id=current_user.id,
        status=status_filter,
    )


@router.get("/saved-jobs/{saved_job_id}", response_model=SavedJobResponse)
def get_saved_job(
    saved_job_id: int,
    current_user: User = Depends(get_current_user),
    repository: SavedJobRepository = Depends(get_saved_job_repository),
):
    saved_job = repository.get_by_id(saved_job_id)

    if saved_job is None:
        raise HTTPException(status_code=404, detail="Saved job not found")

    if saved_job.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")

    return saved_job


@router.post(
    "/saved-jobs/{saved_job_id}/mark-applied",
    response_model=SavedJobResponse,
)
def mark_applied(
    saved_job_id: int,
    current_user: User = Depends(get_current_user),
    repository: SavedJobRepository = Depends(get_saved_job_repository),
):
    saved_job = repository.get_by_id(saved_job_id)

    if saved_job is None:
        raise HTTPException(status_code=404, detail="Saved job not found")

    if saved_job.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")

    return repository.mark_applied(saved_job)


@router.delete(
    "/saved-jobs/{saved_job_id}",
    status_code=status.HTTP_204_NO_CONTENT,
)
def dismiss_saved_job(
    saved_job_id: int,
    current_user: User = Depends(get_current_user),
    repository: SavedJobRepository = Depends(get_saved_job_repository),
):
    saved_job = repository.get_by_id(saved_job_id)

    if saved_job is None:
        raise HTTPException(status_code=404, detail="Saved job not found")

    if saved_job.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")

    repository.delete(saved_job)


# ── Document download endpoints ──────────────────────────────────────


def _make_docx(title: str, body_text: str) -> io.BytesIO:
    """Create a simple .docx from plain text and return as BytesIO."""
    doc = Document()
    doc.add_heading(title, level=1)

    for paragraph in body_text.split("\n"):
        stripped = paragraph.strip()
        if stripped:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@router.get("/saved-jobs/{saved_job_id}/download-resume")
def download_tailored_resume(
    saved_job_id: int,
    current_user: User = Depends(get_current_user),
    repository: SavedJobRepository = Depends(get_saved_job_repository),
):
    """Generate and download a .docx of the tailored resume text."""

    saved_job = repository.get_by_id(saved_job_id)

    if saved_job is None:
        raise HTTPException(status_code=404, detail="Saved job not found")

    if saved_job.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")

    if not saved_job.tailored_resume_text:
        raise HTTPException(
            status_code=404,
            detail="No tailored resume available for this job.",
        )

    company = ""
    if saved_job.job:
        company = saved_job.job.company or "Company"

    buf = _make_docx(
        title=f"Tailored Resume — {company}",
        body_text=saved_job.tailored_resume_text,
    )

    filename = f"tailored_resume_{saved_job_id}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/saved-jobs/{saved_job_id}/download-cover-letter")
def download_cover_letter(
    saved_job_id: int,
    current_user: User = Depends(get_current_user),
    repository: SavedJobRepository = Depends(get_saved_job_repository),
):
    """Generate and download a .docx of the cover letter text."""

    saved_job = repository.get_by_id(saved_job_id)

    if saved_job is None:
        raise HTTPException(status_code=404, detail="Saved job not found")

    if saved_job.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")

    if not saved_job.cover_letter_text:
        raise HTTPException(
            status_code=404,
            detail="No cover letter available for this job.",
        )

    company = ""
    if saved_job.job:
        company = saved_job.job.company or "Company"

    buf = _make_docx(
        title=f"Cover Letter — {company}",
        body_text=saved_job.cover_letter_text,
    )

    filename = f"cover_letter_{saved_job_id}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )