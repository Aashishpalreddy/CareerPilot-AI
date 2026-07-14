from pathlib import Path

from docx import Document
from pypdf import PdfReader


class ResumeParserService:
    """
    Extracts text from supported resume file formats.
    """

    @staticmethod
    def extract_text(file_path: str) -> str:
        extension = Path(file_path).suffix.lower()

        if extension == ".docx":
            return ResumeParserService._parse_docx(file_path)

        if extension == ".pdf":
            return ResumeParserService._parse_pdf(file_path)

        raise ValueError(f"Unsupported file type: {extension}")

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        document = Document(file_path)

        content = []

        # Extract paragraphs
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                content.append(text)

        # Extract tables
        for table in document.tables:
            for row in table.rows:
                row_data = []

                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        row_data.append(text)

                if row_data:
                    content.append(" | ".join(row_data))

        return "\n".join(content)

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)

        content = []

        for page in reader.pages:
            text = page.extract_text()

            if text:
                content.append(text)

        return "\n".join(content)