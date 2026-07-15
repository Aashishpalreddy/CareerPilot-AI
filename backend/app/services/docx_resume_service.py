from pathlib import Path
from uuid import uuid4

from docx import Document

from backend.app.models.resume import Resume
from backend.app.models.tailored_resume import TailoredResume


class DocxResumeService:

    OUTPUT_DIR = Path("generated_resumes")

    @classmethod
    def generate(
        cls,
        resume: Resume,
        tailored_resume: TailoredResume,
    ) -> str:

        cls.OUTPUT_DIR.mkdir(
            exist_ok=True
        )

        document = Document(
            resume.file_path
        )

        paragraphs = document.paragraphs

        summary_updated = False
        experience_updated = False
        projects_updated = False

        i = 0

        while i < len(paragraphs):

            text = paragraphs[i].text.strip().lower()

            # ---------- SUMMARY ----------
            if (
                not summary_updated
                and len(paragraphs[i].text.strip()) > 80
            ):

                paragraphs[i].text = (
                    tailored_resume.tailored_summary or ""
                )

                summary_updated = True

            # ---------- EXPERIENCE ----------
            if (
                text == "experience"
                and not experience_updated
            ):

                i += 1

                while (
                    i < len(paragraphs)
                    and paragraphs[i].text.strip().lower()
                    != "projects"
                ):
                    paragraphs[i].text = ""
                    i += 1

                for job in (
                    tailored_resume.tailored_experience
                    or []
                ):

                    document.add_heading(
                        f"{job.get('title', '')} | {job.get('company', '')}",
                        level=2,
                    )

                    document.add_paragraph(
                        f"{job.get('start_date','')} - {job.get('end_date','')}"
                    )

                    description = job.get(
                        "description",
                        [],
                    )

                    if isinstance(
                        description,
                        list,
                    ):

                        for bullet in description:

                            document.add_paragraph(
                                bullet,
                                style="List Bullet",
                            )

                    else:

                        document.add_paragraph(
                            description
                        )

                experience_updated = True

                continue

            # ---------- PROJECTS ----------
            if (
                text == "projects"
                and not projects_updated
            ):

                i += 1

                while i < len(paragraphs):

                    if (
                        paragraphs[i]
                        .text.strip()
                        .lower()
                        in [
                            "education",
                            "certifications",
                            "skills",
                            "technical skills",
                        ]
                    ):
                        break

                    paragraphs[i].text = ""
                    i += 1

                for project in (
                    tailored_resume.tailored_projects
                    or []
                ):

                    document.add_heading(
                        project.get(
                            "name",
                            "",
                        ),
                        level=2,
                    )

                    document.add_paragraph(
                        project.get(
                            "description",
                            "",
                        )
                    )

                    tech = project.get(
                        "technologies",
                        [],
                    )

                    if tech:

                        document.add_paragraph(
                            "Technologies: "
                            + ", ".join(tech)
                        )

                projects_updated = True

                continue

            i += 1

        output_file = (
            cls.OUTPUT_DIR
            / f"{uuid4()}.docx"
        )

        document.save(
            output_file
        )

        return str(output_file)