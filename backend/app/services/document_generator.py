from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.shared import Pt


class DocumentGenerator:

    def __init__(self):
        self.output_dir = Path("generated_documents/cover_letters")
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_cover_letter(
        self,
        applicant_name: str,
        applicant_email: str,
        company: str,
        position: str,
        content: str,
    ) -> str:

        document = Document()

        heading = document.add_heading(applicant_name, level=1)
        for run in heading.runs:
            run.font.size = Pt(18)

        paragraph = document.add_paragraph()
        paragraph.add_run(applicant_email).font.size = Pt(11)

        document.add_paragraph()

        company_para = document.add_paragraph()
        company_para.add_run("Hiring Manager\n").bold = True
        company_para.add_run(company)

        document.add_paragraph()

        title = document.add_heading(
            f"Application for {position}",
            level=2,
        )

        for run in title.runs:
            run.font.size = Pt(14)

        document.add_paragraph()

        body = document.add_paragraph()

        for line in content.split("\n"):
            body.add_run(line)
            body.add_run("\n")

        document.add_paragraph()

        closing = document.add_paragraph()
        closing.add_run("Sincerely,\n").bold = True
        closing.add_run(applicant_name)

        filename = f"{uuid4()}.docx"
        filepath = self.output_dir / filename

        document.save(filepath)

        return filename